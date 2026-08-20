# -*- coding: utf-8 -*-
"""PaperForge V0 导出:Markdown 全文拼接与 Word(python-docx)生成。"""

import base64
import re
from io import BytesIO

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

CN_FONT = "宋体"
EN_FONT = "Times New Roman"


def to_markdown(payload):
    parts = []
    for c in payload["chapters"]:
        if c["seq"] > 1:
            parts.append("# %s\n\n%s" % (c["title"], c["content_md"]))
        else:
            parts.append(c["content_md"])
    charts = payload.get("chart_suggestions") or []
    if charts:
        lines = ["| 图号 | 图题 | 建议位置 | 所需素材 |", "|---|---|---|---|"]
        for s in charts:
            lines.append("| %s | %s | %s | %s |" % (s["fig"], s["title"], s["position"], s["material"]))
        parts.append("## 图表建议清单\n\n" + "\n".join(lines))
    parts.append("> 本文档由系统自动生成,仅供写作参考,需人工核实后使用。")
    return "\n\n".join(parts)


def _set_font(run, size=12, bold=False, italic=False, color=None):
    run.font.name = EN_FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), EN_FONT)
    rfonts.set(qn("w:hAnsi"), EN_FONT)
    rfonts.set(qn("w:eastAsia"), CN_FONT)


def _strip_bold(text):
    return re.sub(r"\*\*(.+?)\*\*", r"\1", text)


def _add_para(doc, text, style=None, align=None, size=12, bold=False):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.4
    _set_font(p.add_run(_strip_bold(text)), size=size, bold=bold)
    return p


def _add_md_table(doc, rows):
    header = [c.strip() for c in rows[0].strip().strip("|").split("|")]
    body = [
        [c.strip() for c in row.strip().strip("|").split("|")]
        for row in rows[1:]
        if row.strip() and not re.match(r"^\s*\|[\s:\-|]+\|\s*$", row)
    ]
    if not body:
        return
    ncols = len(header)
    table = doc.add_table(rows=1 + len(body), cols=ncols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(header):
        cell = table.rows[0].cells[j]
        cell.paragraphs[0].paragraph_format.space_after = Pt(2)
        _set_font(cell.paragraphs[0].add_run(h), size=10.5, bold=True)
    for i, row in enumerate(body, start=1):
        for j in range(ncols):
            value = row[j] if j < len(row) else ""
            cell = table.rows[i].cells[j]
            cell.paragraphs[0].paragraph_format.space_after = Pt(2)
            _set_font(cell.paragraphs[0].add_run(value), size=10.5)
    for row in table.rows:
        for cell in row.cells:
            tcpr = cell._tc.get_or_add_tcPr()
            v = OxmlElement("w:vAlign")
            v.set(qn("w:val"), "center")
            tcpr.append(v)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def _render_md(doc, content):
    lines = content.splitlines()
    i = 0
    in_code = False
    while i < len(lines):
        line = lines[i].rstrip()
        img_match = re.match(r"^!\[(.*?)\]\((data:image/png;base64,[A-Za-z0-9+/=]+)\)\s*$", line)
        if img_match:
            caption = img_match.group(1)
            b64 = img_match.group(2).split(",", 1)[1]
            try:
                doc.add_picture(BytesIO(base64.b64decode(b64)), width=Inches(5.6))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                _add_para(doc, caption, align=WD_ALIGN_PARAGRAPH.CENTER, size=10.5, italic=True)
            except Exception:
                _add_para(doc, line, size=10)
            i += 1
            continue
        if line.startswith("```"):
            in_code = not in_code
            i += 1
            continue
        if in_code:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            _set_font(p.add_run(line), size=10, color=RGBColor(0x33, 0x33, 0x33))
            i += 1
            continue
        if line.startswith("|"):
            rows = [line]
            j = i + 1
            while j < len(lines) and lines[j].strip().startswith("|"):
                rows.append(lines[j].strip())
                j += 1
            _add_md_table(doc, rows)
            i = j
            continue
        if line.startswith("### "):
            _add_para(doc, line[4:], style="Heading 3", size=13, bold=True)
        elif line.startswith("## "):
            _add_para(doc, line[3:], style="Heading 2", size=14, bold=True)
        elif line.startswith("# "):
            _add_para(doc, line[2:], style="Heading 1", size=16, bold=True)
        elif re.match(r"^\s*[-*]\s+", line):
            _add_para(doc, re.sub(r"^\s*[-*]\s+", "", line), style="List Bullet", size=12)
        elif re.match(r"^\s*\d+[.、]\s+", line):
            _add_para(doc, re.sub(r"^\s*\d+[.、]\s+", "", line), style="List Number", size=12)
        elif line.startswith(">"):
            _add_para(doc, line.lstrip("> ").strip(), style="Intense Quote", size=11, italic=True)
        elif line.strip():
            _add_para(doc, line, size=12)
        i += 1


def build_docx_bytes(payload):
    doc = Document()
    for section in doc.sections:
        section.left_margin = Inches(1.1)
        section.right_margin = Inches(1.1)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)

    title = payload.get("title", "论文初稿")
    _add_para(doc, title, align=WD_ALIGN_PARAGRAPH.CENTER, size=22, bold=True)
    _add_para(
        doc,
        "AI 辅助论文初稿 · 仅供参考 · 生成时间 %s" % payload.get("generated_at", ""),
        align=WD_ALIGN_PARAGRAPH.CENTER,
        size=10.5,
    )
    doc.add_paragraph()

    for chapter in payload["chapters"]:
        if chapter["seq"] > 1:
            _add_para(doc, chapter["title"], style="Heading 1", size=16, bold=True)
        _render_md(doc, chapter["content_md"])

    charts = payload.get("chart_suggestions") or []
    if charts:
        doc.add_paragraph()
        _add_para(doc, "图表建议清单", style="Heading 1", size=16, bold=True)
        _add_md_table(
            doc,
            ["| 图号 | 图题 | 建议位置 | 所需素材 |"]
            + [
                "| %s | %s | %s | %s |" % (s["fig"], s["title"], s["position"], s["material"])
                for s in charts
            ],
        )

    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(12)
    _set_font(note.add_run("声明:本文档由系统自动生成,仅供写作参考,需人工核实后使用;封面作者信息请自行填写。"), size=10.5, italic=True)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
